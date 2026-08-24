"""Turn a downloaded submission file into plain text.

Every function returns (text, note). `note` is non-empty when the text is
unreliable or missing — the runner surfaces those to the teacher instead of
drafting feedback on a bad read.
"""
from pathlib import Path

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tif", ".tiff", ".webp"}
AUDIO_EXTS = {".mp3", ".m4a", ".ogg", ".oga", ".opus", ".wav", ".aac", ".amr", ".webm"}
DOC_EXTS = {".docx"}
PDF_EXTS = {".pdf"}

# Below this many characters a "successful" extraction is almost certainly a
# failed one — a blank scan, or OCR that found nothing but noise.
MIN_USEFUL_CHARS = 40

_whisper_model = None


def kind_of(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in IMAGE_EXTS:
        return "image"
    if ext in AUDIO_EXTS:
        return "audio"
    if ext in DOC_EXTS:
        return "docx"
    if ext in PDF_EXTS:
        return "pdf"
    return "unknown"


def from_docx(path: Path) -> tuple[str, str]:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells if c.text.strip())
    text = "\n".join(parts).strip()
    if not text:
        return "", "the Word document had no readable text in it"
    return text, ""


def from_pdf(path: Path) -> tuple[str, str]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if len(text) >= MIN_USEFUL_CHARS:
        return text, ""
    # No text layer: it's a scan, so fall through to OCR of the rendered pages.
    try:
        from pdf2image import convert_from_path
    except ImportError:
        return text, "the PDF looks scanned and pdf2image is not installed to OCR it"
    try:
        pages = convert_from_path(str(path), dpi=300)
    except Exception as exc:  # poppler missing, encrypted PDF, etc.
        return text, f"the PDF looks scanned and could not be rendered for OCR ({exc})"
    ocr = "\n".join(_ocr_image(page) for page in pages).strip()
    if len(ocr) < MIN_USEFUL_CHARS:
        return ocr, "OCR of the scanned PDF came back nearly empty"
    return ocr, ""


def from_image(path: Path) -> tuple[str, str]:
    from PIL import Image

    text = _ocr_image(Image.open(path)).strip()
    if len(text) < MIN_USEFUL_CHARS:
        return text, "OCR found little or no text (handwriting or a low-quality photo)"
    return text, ""


def _ocr_image(image) -> str:
    import pytesseract

    return pytesseract.image_to_string(image)


def from_audio(path: Path, model_size: str, language: str | None) -> tuple[str, str]:
    global _whisper_model
    try:
        from faster_whisper import WhisperModel
    except ImportError:
        return "", "faster-whisper is not installed, so the voice message was not transcribed"
    if _whisper_model is None:
        print(f"  (loading the {model_size} transcription model, first run takes a minute)")
        _whisper_model = WhisperModel(model_size, device="cpu", compute_type="int8")
    segments, _info = _whisper_model.transcribe(str(path), language=language)
    text = " ".join(s.text.strip() for s in segments).strip()
    if len(text) < MIN_USEFUL_CHARS:
        return text, "the transcription came back nearly empty (quiet or very short recording)"
    return text, ""


def extract(path: Path, config: dict) -> tuple[str, str]:
    kind = kind_of(path)
    transcription = config.get("transcription", {})
    try:
        if kind == "docx":
            return from_docx(path)
        if kind == "pdf":
            return from_pdf(path)
        if kind == "image":
            return from_image(path)
        if kind == "audio":
            return from_audio(
                path,
                transcription.get("model", "small"),
                transcription.get("language"),
            )
    except Exception as exc:
        return "", f"reading the file failed: {exc}"
    return "", f"'{path.suffix}' is not a file type this script knows how to read"
